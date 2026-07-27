"""
export_utils.py
===============
Render a plan (plus its computed budget) into Markdown, plain text, DOCX,
CSV and JSON.

Every exporter consumes the same normalised plan dict plus a ``BudgetSummary``
from ``budget.py``, so no exporter ever performs its own arithmetic.
"""

from __future__ import annotations

import csv
import io
import json
from datetime import datetime, timedelta
from typing import Any

from budget import BudgetSummary


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #


def fmt_days(days: int) -> str:
    """Turn a relative day offset into human wording."""
    if days == 0:
        return "Event day"
    if days < 0:
        n = abs(days)
        return f"{n} day{'s' if n != 1 else ''} after"
    if days >= 7 and days % 7 == 0:
        weeks = days // 7
        return f"{weeks} week{'s' if weeks != 1 else ''} before"
    return f"{days} day{'s' if days != 1 else ''} before"


def actual_date(event_date: str, days_before: int) -> str:
    """Convert a relative offset into a real date, if the event date is known."""
    if not event_date:
        return ""
    for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y", "%d %B %Y", "%B %d, %Y"):
        try:
            base = datetime.strptime(event_date.strip(), fmt)
            return (base - timedelta(days=days_before)).strftime("%d %b %Y")
        except ValueError:
            continue
    return ""


def _money(amount: float, symbol: str) -> str:
    return f"{symbol}{amount:,.0f}"


def _wrap(text: str, width: int) -> list[str]:
    words = str(text).split()
    if not words:
        return [""]
    lines: list[str] = []
    current = ""
    for word in words:
        if len(current) + len(word) + 1 <= width:
            current = f"{current} {word}".strip()
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


# --------------------------------------------------------------------------- #
# Markdown
# --------------------------------------------------------------------------- #


def to_markdown(
    plan: dict[str, Any],
    bud: BudgetSummary,
    symbol: str = "₹",
    event_date: str = "",
    attendees: int = 0,
) -> str:
    L: list[str] = []
    add = L.append

    add(f"# {plan.get('event_title', 'Event Plan')}")
    add("")
    if event_date:
        add(f"**Date:** {event_date}  ")
    if attendees:
        add(f"**Expected attendees:** {attendees}  ")
    add(f"**Estimated cost:** {_money(bud.grand_total, symbol)}  ")
    if attendees:
        add(f"**Per head:** {_money(bud.per_head, symbol)}  ")
    add(f"**Generated:** {datetime.now().strftime('%d %B %Y, %H:%M')}  ")
    add("")
    add("---")
    add("")

    if plan.get("summary"):
        add("## Overview")
        add("")
        add(plan["summary"])
        add("")

    # Venue
    venue = plan.get("venue", {})
    if venue.get("recommendations"):
        add("## Venue Recommendations")
        add("")
        for rec in venue["recommendations"]:
            add(f"### {rec.get('option', '')}")
            if rec.get("why"):
                add(f"{rec['why']}")
            if rec.get("capacity_fit"):
                add(f"- **Capacity:** {rec['capacity_fit']}")
            try:
                cost = float(rec.get("est_cost") or 0)
                if cost > 0:
                    add(f"- **Estimated cost:** {_money(cost, symbol)}")
            except (TypeError, ValueError):
                pass
            add("")
    if venue.get("layout"):
        add("**Recommended layout:** " + venue["layout"])
        add("")
    if venue.get("requirements"):
        add("**Venue must have:**")
        add("")
        for req in venue["requirements"]:
            add(f"- {req}")
        add("")

    # Budget
    if bud.lines:
        add("## Budget")
        add("")
        add("| Category | Item | Unit Cost | Qty | Total |")
        add("|---|---|---:|---:|---:|")
        for line in bud.lines:
            item = line.item.replace("|", "\\|")
            add(
                f"| {line.category} | {item} | {_money(line.unit_cost, symbol)} | "
                f"{line.quantity:g} {line.unit} | {_money(line.total, symbol)} |"
            )
        add("")
        add(f"**Subtotal:** {_money(bud.subtotal, symbol)}  ")
        add(f"**Contingency ({bud.contingency_pct:.0f}%):** {_money(bud.contingency, symbol)}  ")
        add(f"**Grand total:** {_money(bud.grand_total, symbol)}  ")
        if bud.target > 0:
            state = "over" if bud.over_budget else "under"
            add(f"**Against a budget of {_money(bud.target, symbol)}:** "
                f"{_money(abs(bud.variance), symbol)} {state} ({abs(bud.variance_pct):.1f}%)  ")
        add("")
        add("### Spend by category")
        add("")
        for cat, amount in sorted(bud.by_category.items(), key=lambda kv: -kv[1]):
            pct = bud.category_pct.get(cat, 0)
            add(f"- **{cat}:** {_money(amount, symbol)} ({pct:.1f}%)")
        add("")

    # Checklist
    if plan.get("checklist"):
        add("## Checklist")
        add("")
        current_phase = None
        for task in plan["checklist"]:
            if task["phase"] != current_phase:
                current_phase = task["phase"]
                add(f"### {current_phase}")
                add("")
            when = fmt_days(task["days_before_event"])
            date_str = actual_date(event_date, task["days_before_event"])
            when_full = f"{when} ({date_str})" if date_str else when
            add(f"- [ ] **{task['task']}** — {task['owner_role']} · {when_full} · {task['priority']}")
        add("")

    # Timeline
    if plan.get("timeline"):
        add("## Timeline")
        add("")
        for ms in plan["timeline"]:
            when = fmt_days(ms["days_before_event"])
            date_str = actual_date(event_date, ms["days_before_event"])
            when_full = f"{when} ({date_str})" if date_str else when
            add(f"- **{when_full}** — {ms['milestone']}: {ms.get('detail', '')}")
        add("")

    # Equipment
    if plan.get("equipment"):
        add("## Equipment")
        add("")
        add("| Item | Quantity | Essential | Notes |")
        add("|---|---|---|---|")
        for eq in plan["equipment"]:
            add(f"| {eq['item']} | {eq['quantity']} | "
                f"{'Yes' if eq['essential'] else 'Optional'} | {eq.get('notes', '')} |")
        add("")

    # Invitations
    inv = plan.get("invitations", {})
    if inv.get("channels") or inv.get("sample_message"):
        add("## Invitations")
        add("")
        if inv.get("channels"):
            add(f"**Channels:** {', '.join(inv['channels'])}  ")
        if inv.get("send_schedule"):
            add(f"**Schedule:** {inv['send_schedule']}  ")
        if inv.get("rsvp_method"):
            add(f"**RSVP:** {inv['rsvp_method']}  ")
        add("")
        if inv.get("guest_segments"):
            add("| Segment | Count | Approach |")
            add("|---|---:|---|")
            for seg in inv["guest_segments"]:
                add(f"| {seg['segment']} | {seg['count']} | {seg.get('approach', '')} |")
            add("")
        if inv.get("sample_message"):
            add("**Sample invitation:**")
            add("")
            add("> " + inv["sample_message"].replace("\n", "\n> "))
            add("")

    # Day schedule
    if plan.get("day_schedule"):
        add("## Event Day Schedule")
        add("")
        add("| Time | Activity | Owner |")
        add("|---|---|---|")
        for slot in plan["day_schedule"]:
            add(f"| {slot.get('time', '')} | {slot['activity']} | {slot.get('owner', '')} |")
        add("")

    # Risks
    if plan.get("risks"):
        add("## Risks & Mitigation")
        add("")
        for risk in plan["risks"]:
            add(f"- **{risk['risk']}** *(likelihood: {risk['likelihood']})*  ")
            if risk.get("mitigation"):
                add(f"  → {risk['mitigation']}")
        add("")

    if plan.get("tips"):
        add("## Tips")
        add("")
        for tip in plan["tips"]:
            add(f"- {tip}")
        add("")

    add("---")
    add("")
    add("*Generated by AI Event Planning Agent*")
    return "\n".join(L)


# --------------------------------------------------------------------------- #
# Plain text
# --------------------------------------------------------------------------- #


def to_text(
    plan: dict[str, Any],
    bud: BudgetSummary,
    symbol: str = "₹",
    event_date: str = "",
    attendees: int = 0,
) -> str:
    W = 76
    L: list[str] = []
    add = L.append

    add("=" * W)
    add(str(plan.get("event_title", "Event Plan")).upper().center(W))
    add("=" * W)
    add("")
    if event_date:
        add(f"Date            : {event_date}")
    if attendees:
        add(f"Attendees       : {attendees}")
    add(f"Estimated cost  : {_money(bud.grand_total, symbol)}")
    if attendees:
        add(f"Per head        : {_money(bud.per_head, symbol)}")
    add(f"Generated       : {datetime.now().strftime('%d %B %Y, %H:%M')}")

    def section(name: str) -> None:
        add("")
        add(name.upper())
        add("-" * W)

    if plan.get("summary"):
        section("Overview")
        for line in _wrap(plan["summary"], W):
            add(line)

    if bud.lines:
        section("Budget")
        for line in bud.lines:
            add(f"  {line.item[:44]:<44} {_money(line.total, symbol):>14}")
            add(f"    {line.category} · {_money(line.unit_cost, symbol)} x {line.quantity:g} {line.unit}")
        add("  " + "-" * (W - 4))
        add(f"  {'Subtotal':<44} {_money(bud.subtotal, symbol):>14}")
        add(f"  {f'Contingency ({bud.contingency_pct:.0f}%)':<44} {_money(bud.contingency, symbol):>14}")
        add(f"  {'GRAND TOTAL':<44} {_money(bud.grand_total, symbol):>14}")

    if plan.get("checklist"):
        section("Checklist")
        current_phase = None
        for task in plan["checklist"]:
            if task["phase"] != current_phase:
                current_phase = task["phase"]
                add("")
                add(f"  [{current_phase}]")
            when = fmt_days(task["days_before_event"])
            date_str = actual_date(event_date, task["days_before_event"])
            for j, line in enumerate(_wrap(task["task"], W - 8)):
                add(("    [ ] " if j == 0 else "        ") + line)
            add(f"        {task['owner_role']} · {when}"
                + (f" ({date_str})" if date_str else "")
                + f" · {task['priority']}")

    if plan.get("timeline"):
        section("Timeline")
        for ms in plan["timeline"]:
            when = fmt_days(ms["days_before_event"])
            date_str = actual_date(event_date, ms["days_before_event"])
            add(f"  {when}{' (' + date_str + ')' if date_str else ''}")
            add(f"    {ms['milestone']}")
            if ms.get("detail"):
                for line in _wrap(ms["detail"], W - 6):
                    add("      " + line)

    if plan.get("equipment"):
        section("Equipment")
        for eq in plan["equipment"]:
            mark = "*" if eq["essential"] else " "
            add(f"  {mark} {eq['item'][:50]:<50} {eq['quantity']}")
        add("")
        add("  * = essential")

    inv = plan.get("invitations", {})
    if inv.get("sample_message"):
        section("Sample Invitation")
        for line in _wrap(inv["sample_message"], W - 4):
            add("  " + line)

    if plan.get("day_schedule"):
        section("Event Day Schedule")
        for slot in plan["day_schedule"]:
            add(f"  {slot.get('time', ''):<8} {slot['activity'][:52]:<52} {slot.get('owner', '')}")

    if plan.get("risks"):
        section("Risks & Mitigation")
        for risk in plan["risks"]:
            for j, line in enumerate(_wrap(risk["risk"], W - 8)):
                add(("  ! " if j == 0 else "    ") + line)
            add(f"      Likelihood: {risk['likelihood']}")
            if risk.get("mitigation"):
                for line in _wrap("Mitigation: " + risk["mitigation"], W - 8):
                    add("      " + line)

    if plan.get("tips"):
        section("Tips")
        for tip in plan["tips"]:
            for j, line in enumerate(_wrap(tip, W - 6)):
                add(("  - " if j == 0 else "    ") + line)

    add("")
    add("=" * W)
    add("Generated by AI Event Planning Agent".center(W))
    add("=" * W)
    return "\n".join(L)


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #


def to_docx(
    plan: dict[str, Any],
    bud: BudgetSummary,
    symbol: str = "₹",
    event_date: str = "",
    attendees: int = 0,
) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    INK = RGBColor(0x1A, 0x14, 0x10)
    CORAL = RGBColor(0xF2, 0x54, 0x2D)
    STONE = RGBColor(0x8A, 0x75, 0x63)
    PLUM = RGBColor(0x7B, 0x2D, 0x5E)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(str(plan.get("event_title", "Event Plan")))
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = INK

    meta_bits = []
    if event_date:
        meta_bits.append(event_date)
    if attendees:
        meta_bits.append(f"{attendees} attendees")
    meta_bits.append(f"Est. {_money(bud.grand_total, symbol)}")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run("  •  ".join(meta_bits))
    srun.font.size = Pt(9.5)
    srun.font.color.rgb = STONE

    doc.add_paragraph("─" * 58).alignment = WD_ALIGN_PARAGRAPH.CENTER

    def heading(text: str, color=CORAL) -> None:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text.upper())
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = color

    if plan.get("summary"):
        heading("Overview")
        doc.add_paragraph(plan["summary"])

    venue = plan.get("venue", {})
    if venue.get("recommendations"):
        heading("Venue Recommendations")
        for rec in venue["recommendations"]:
            para = doc.add_paragraph()
            run = para.add_run(str(rec.get("option", "")))
            run.font.bold = True
            run.font.size = Pt(11)
            try:
                cost = float(rec.get("est_cost") or 0)
                if cost > 0:
                    crun = para.add_run(f"   {_money(cost, symbol)}")
                    crun.font.size = Pt(9.5)
                    crun.font.color.rgb = STONE
            except (TypeError, ValueError):
                pass
            if rec.get("why"):
                doc.add_paragraph(str(rec["why"]))
            if rec.get("capacity_fit"):
                cap = doc.add_paragraph(f"Capacity: {rec['capacity_fit']}")
                cap.runs[0].font.size = Pt(9)
                cap.runs[0].font.color.rgb = STONE
    if venue.get("layout"):
        heading("Recommended Layout")
        doc.add_paragraph(venue["layout"])
    if venue.get("requirements"):
        heading("Venue Requirements")
        for req in venue["requirements"]:
            doc.add_paragraph(str(req), style="List Bullet")

    if bud.lines:
        heading("Budget")
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 2"
        headers = table.rows[0].cells
        for idx, label in enumerate(["Category", "Item", "Unit", "Qty", "Total"]):
            headers[idx].text = label
            for para in headers[idx].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        for line in bud.lines:
            cells = table.add_row().cells
            cells[0].text = line.category
            cells[1].text = line.item
            cells[2].text = _money(line.unit_cost, symbol)
            cells[3].text = f"{line.quantity:g} {line.unit}"
            cells[4].text = _money(line.total, symbol)
            for cell in cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

        for label, value, bold in [
            ("Subtotal", bud.subtotal, False),
            (f"Contingency ({bud.contingency_pct:.0f}%)", bud.contingency, False),
            ("Grand Total", bud.grand_total, True),
        ]:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run(f"{label}:  {_money(value, symbol)}")
            run.font.bold = bold
            run.font.size = Pt(11 if bold else 10)
            if bold:
                run.font.color.rgb = CORAL

    if plan.get("checklist"):
        heading("Checklist")
        current_phase = None
        for task in plan["checklist"]:
            if task["phase"] != current_phase:
                current_phase = task["phase"]
                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(9)
                run = para.add_run(current_phase)
                run.font.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = PLUM
            when = fmt_days(task["days_before_event"])
            date_str = actual_date(event_date, task["days_before_event"])
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(f"{task['task']}").font.size = Pt(10)
            detail = para.add_run(
                f"  —  {task['owner_role']} · {when}"
                + (f" ({date_str})" if date_str else "")
                + f" · {task['priority']}"
            )
            detail.font.size = Pt(8.5)
            detail.font.color.rgb = STONE

    if plan.get("timeline"):
        heading("Timeline")
        for ms in plan["timeline"]:
            when = fmt_days(ms["days_before_event"])
            date_str = actual_date(event_date, ms["days_before_event"])
            para = doc.add_paragraph()
            run = para.add_run(f"{when}{' (' + date_str + ')' if date_str else ''}  ")
            run.font.bold = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = CORAL
            para.add_run(f"{ms['milestone']}").font.size = Pt(10)
            if ms.get("detail"):
                sub = doc.add_paragraph(str(ms["detail"]))
                sub.paragraph_format.left_indent = Pt(18)
                sub.runs[0].font.size = Pt(9)
                sub.runs[0].font.color.rgb = STONE

    if plan.get("equipment"):
        heading("Equipment")
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 2"
        headers = table.rows[0].cells
        for idx, label in enumerate(["Item", "Qty", "Essential", "Notes"]):
            headers[idx].text = label
            for para in headers[idx].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        for eq in plan["equipment"]:
            cells = table.add_row().cells
            cells[0].text = str(eq["item"])
            cells[1].text = str(eq["quantity"])
            cells[2].text = "Yes" if eq["essential"] else "Optional"
            cells[3].text = str(eq.get("notes", ""))
            for cell in cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

    inv = plan.get("invitations", {})
    if inv.get("channels") or inv.get("sample_message"):
        heading("Invitations")
        if inv.get("channels"):
            doc.add_paragraph(f"Channels: {', '.join(inv['channels'])}")
        if inv.get("send_schedule"):
            doc.add_paragraph(f"Schedule: {inv['send_schedule']}")
        if inv.get("rsvp_method"):
            doc.add_paragraph(f"RSVP: {inv['rsvp_method']}")
        if inv.get("sample_message"):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(18)
            para.paragraph_format.space_before = Pt(8)
            run = para.add_run(str(inv["sample_message"]))
            run.font.italic = True
            run.font.size = Pt(10)

    if plan.get("day_schedule"):
        heading("Event Day Schedule")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 2"
        headers = table.rows[0].cells
        for idx, label in enumerate(["Time", "Activity", "Owner"]):
            headers[idx].text = label
            for para in headers[idx].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
        for slot in plan["day_schedule"]:
            cells = table.add_row().cells
            cells[0].text = str(slot.get("time", ""))
            cells[1].text = str(slot["activity"])
            cells[2].text = str(slot.get("owner", ""))
            for cell in cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)

    if plan.get("risks"):
        heading("Risks & Mitigation")
        for risk in plan["risks"]:
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(str(risk["risk"]))
            run.font.bold = True
            run.font.size = Pt(10)
            lik = para.add_run(f"  ({risk['likelihood']})")
            lik.font.size = Pt(8.5)
            lik.font.color.rgb = STONE
            if risk.get("mitigation"):
                sub = doc.add_paragraph(str(risk["mitigation"]))
                sub.paragraph_format.left_indent = Pt(28)
                sub.runs[0].font.size = Pt(9)
                sub.runs[0].font.color.rgb = STONE

    if plan.get("tips"):
        heading("Tips")
        for tip in plan["tips"]:
            doc.add_paragraph(str(tip), style="List Bullet")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(20)
    frun = footer.add_run(
        f"Generated by AI Event Planning Agent  •  {datetime.now().strftime('%d %B %Y, %H:%M')}"
    )
    frun.font.size = Pt(8)
    frun.font.color.rgb = STONE

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# --------------------------------------------------------------------------- #
# CSV / JSON
# --------------------------------------------------------------------------- #


def checklist_to_csv(plan: dict[str, Any], event_date: str = "") -> str:
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(["#", "Phase", "Task", "Owner", "When", "Due Date", "Priority", "Done"])
    for i, task in enumerate(plan.get("checklist", []), 1):
        writer.writerow(
            [
                i,
                task["phase"],
                task["task"],
                task["owner_role"],
                fmt_days(task["days_before_event"]),
                actual_date(event_date, task["days_before_event"]),
                task["priority"],
                "",
            ]
        )
    return buffer.getvalue()


def budget_to_csv(bud: BudgetSummary) -> str:
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(["#", "Category", "Item", "Unit Cost", "Quantity", "Unit", "Total", "Notes"])
    for i, line in enumerate(bud.lines, 1):
        writer.writerow(
            [i, line.category, line.item, line.unit_cost, line.quantity, line.unit, line.total, line.notes]
        )
    writer.writerow([])
    writer.writerow(["", "", "Subtotal", "", "", "", bud.subtotal, ""])
    writer.writerow(["", "", f"Contingency ({bud.contingency_pct:.0f}%)", "", "", "", bud.contingency, ""])
    writer.writerow(["", "", "GRAND TOTAL", "", "", "", bud.grand_total, ""])
    return buffer.getvalue()


def to_json(plan: dict[str, Any], bud: BudgetSummary) -> str:
    payload = dict(plan)
    payload["computed_budget"] = {
        "lines": [
            {
                "category": line.category,
                "item": line.item,
                "unit_cost": line.unit_cost,
                "quantity": line.quantity,
                "unit": line.unit,
                "total": line.total,
                "notes": line.notes,
            }
            for line in bud.lines
        ],
        "subtotal": bud.subtotal,
        "contingency_pct": bud.contingency_pct,
        "contingency": bud.contingency,
        "grand_total": bud.grand_total,
        "per_head": bud.per_head,
        "by_category": bud.by_category,
        "category_pct": bud.category_pct,
        "target": bud.target,
        "variance": bud.variance,
    }
    return json.dumps(payload, indent=2, ensure_ascii=False)


def safe_filename(title: str, extension: str) -> str:
    base = "".join(ch if ch.isalnum() or ch in " -_" else "" for ch in str(title))
    base = "_".join(base.split())[:50] or "event_plan"
    return f"{base}_{datetime.now().strftime('%Y%m%d')}.{extension}"
