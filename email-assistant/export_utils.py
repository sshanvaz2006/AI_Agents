"""
export_utils.py
----------------
Helpers to package a generated email as a downloadable .docx file.
"""

import io
from docx import Document
from docx.shared import Pt


def build_docx_bytes(subject: str, body: str) -> bytes:
    doc = Document()

    title = doc.add_heading(level=1)
    run = title.add_run(subject)
    run.font.size = Pt(16)

    doc.add_paragraph()  # spacing

    for para in body.split("\n"):
        p = doc.add_paragraph(para)
        for r in p.runs:
            r.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def build_txt_bytes(subject: str, body: str) -> bytes:
    content = f"Subject: {subject}\n\n{body}"
    return content.encode("utf-8")
